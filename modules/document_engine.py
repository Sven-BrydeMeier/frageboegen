"""
Dokument-Engine
Generiert DOCX und PDF Dokumente aus Templates mit Platzhaltern
"""

import os
import io
import re
from typing import Dict, Any, Optional, List, Union
from datetime import datetime, date
from pathlib import Path
import tempfile
import hashlib

# DOCX
try:
    from docxtpl import DocxTemplate
    HAS_DOCXTPL = True
except ImportError:
    HAS_DOCXTPL = False

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# PDF
# WeasyPrint braucht System-Bibliotheken (libpango) - auf Streamlit Cloud nicht verfügbar
HAS_WEASYPRINT = False
_weasyprint_error = None

def _try_import_weasyprint():
    """Lazy import für WeasyPrint"""
    global HAS_WEASYPRINT, _weasyprint_error
    try:
        from weasyprint import HTML, CSS
        HAS_WEASYPRINT = True
        return HTML, CSS
    except (ImportError, OSError) as e:
        _weasyprint_error = str(e)
        HAS_WEASYPRINT = False
        return None, None

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm, mm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
    from reportlab.pdfgen import canvas
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

try:
    from pypdf import PdfReader, PdfWriter, PdfMerger
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

from jinja2 import Environment, BaseLoader, FileSystemLoader


class DocumentEngine:
    """Haupt-Engine für Dokumentengenerierung"""
    
    def __init__(self, template_dir: Optional[str] = None):
        self.template_dir = Path(template_dir) if template_dir else Path("templates")
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.template_dir)) if self.template_dir.exists() else BaseLoader()
        )
        self._setup_jinja_filters()
    
    def _setup_jinja_filters(self):
        """Registriert benutzerdefinierte Jinja2-Filter"""
        
        def format_date(value, format="%d.%m.%Y"):
            """Formatiert ein Datum"""
            if isinstance(value, str):
                try:
                    value = datetime.fromisoformat(value)
                except:
                    return value
            if isinstance(value, (datetime, date)):
                return value.strftime(format)
            return value
        
        def format_currency(value, symbol="€", decimals=2):
            """Formatiert als Währung"""
            try:
                return f"{float(value):,.{decimals}f} {symbol}".replace(",", "X").replace(".", ",").replace("X", ".")
            except:
                return value
        
        def format_phone(value):
            """Formatiert Telefonnummer"""
            if not value:
                return ""
            # Nur Ziffern behalten
            digits = re.sub(r'\D', '', str(value))
            if len(digits) >= 10:
                # Deutsche Formatierung
                return f"{digits[:4]} {digits[4:7]} {digits[7:]}"
            return value
        
        def yes_no(value, yes="Ja", no="Nein"):
            """Konvertiert Boolean zu Ja/Nein"""
            if isinstance(value, bool):
                return yes if value else no
            if str(value).lower() in ['true', '1', 'ja', 'yes']:
                return yes
            return no
        
        def default_if_empty(value, default="—"):
            """Gibt Default zurück wenn leer"""
            if value is None or value == "" or value == []:
                return default
            return value
        
        def join_list(value, separator=", "):
            """Verbindet Liste zu String"""
            if isinstance(value, list):
                return separator.join(str(v) for v in value)
            return value
        
        self.jinja_env.filters['date'] = format_date
        self.jinja_env.filters['currency'] = format_currency
        self.jinja_env.filters['phone'] = format_phone
        self.jinja_env.filters['yesno'] = yes_no
        self.jinja_env.filters['default_if_empty'] = default_if_empty
        self.jinja_env.filters['join'] = join_list
    
    def render_template_string(self, template_str: str, data: Dict[str, Any]) -> str:
        """Rendert einen Template-String mit Daten"""
        template = self.jinja_env.from_string(template_str)
        return template.render(**data)
    
    # ========================================
    # DOCX Generierung
    # ========================================
    
    def generate_docx_from_template(
        self,
        template_path: Union[str, Path],
        data: Dict[str, Any],
        output_path: Optional[Union[str, Path]] = None
    ) -> bytes:
        """
        Generiert DOCX aus Word-Template mit Jinja2-Platzhaltern
        
        Template-Syntax in Word:
        - {{ variable }} - Einfacher Platzhalter
        - {% if bedingung %}...{% endif %} - Bedingung
        - {% for item in liste %}...{% endfor %} - Schleife
        - {{ wert|date }} - Filter anwenden
        """
        if not HAS_DOCXTPL:
            raise ImportError("docxtpl ist nicht installiert: pip install docxtpl")
        
        template_path = Path(template_path)
        if not template_path.is_absolute():
            template_path = self.template_dir / template_path
        
        if not template_path.exists():
            raise FileNotFoundError(f"Template nicht gefunden: {template_path}")
        
        # Template laden und rendern
        doc = DocxTemplate(template_path)
        
        # Jinja-Filter registrieren
        doc.jinja_env.filters.update(self.jinja_env.filters)
        
        # Kontext vorbereiten (Datum etc.)
        context = {
            **data,
            'heute': datetime.now(),
            'datum': datetime.now().strftime("%d.%m.%Y"),
            'uhrzeit': datetime.now().strftime("%H:%M"),
        }
        
        # Rendern
        doc.render(context)
        
        # Speichern
        if output_path:
            doc.save(output_path)
            with open(output_path, 'rb') as f:
                return f.read()
        else:
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.read()
    
    def generate_docx_programmatic(
        self,
        title: str,
        data: Dict[str, Any],
        sections: List[Dict[str, Any]],
        output_path: Optional[Union[str, Path]] = None
    ) -> bytes:
        """
        Generiert DOCX programmatisch (ohne Template)
        
        sections = [
            {"type": "heading", "level": 1, "text": "Titel"},
            {"type": "paragraph", "text": "Text..."},
            {"type": "table", "headers": [...], "rows": [...]},
            {"type": "list", "items": [...]},
        ]
        """
        if not HAS_DOCX:
            raise ImportError("python-docx ist nicht installiert: pip install python-docx")
        
        doc = Document()
        
        # Titel
        doc.add_heading(title, 0)
        doc.add_paragraph(f"Erstellt am {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        doc.add_paragraph()
        
        # Sections verarbeiten
        for section in sections:
            section_type = section.get("type", "paragraph")
            
            if section_type == "heading":
                level = section.get("level", 1)
                text = self.render_template_string(section.get("text", ""), data)
                doc.add_heading(text, level)
            
            elif section_type == "paragraph":
                text = self.render_template_string(section.get("text", ""), data)
                p = doc.add_paragraph(text)
                if section.get("bold"):
                    p.runs[0].bold = True
            
            elif section_type == "table":
                headers = section.get("headers", [])
                rows = section.get("rows", [])
                
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'
                
                # Header
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                
                # Rows
                for row_data in rows:
                    row = table.add_row().cells
                    for i, cell_data in enumerate(row_data):
                        if i < len(row):
                            text = self.render_template_string(str(cell_data), data)
                            row[i].text = text
            
            elif section_type == "list":
                items = section.get("items", [])
                for item in items:
                    text = self.render_template_string(str(item), data)
                    doc.add_paragraph(text, style='List Bullet')
            
            elif section_type == "field_value":
                label = section.get("label", "")
                field = section.get("field", "")
                value = data.get(field, "—")
                if isinstance(value, list):
                    value = ", ".join(str(v) for v in value)
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(value))
            
            elif section_type == "spacer":
                doc.add_paragraph()
        
        # Speichern
        if output_path:
            doc.save(output_path)
            with open(output_path, 'rb') as f:
                return f.read()
        else:
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.read()
    
    # ========================================
    # PDF Generierung
    # ========================================
    
    def generate_pdf_from_html(
        self,
        html_content: str,
        data: Dict[str, Any],
        css: Optional[str] = None,
        output_path: Optional[Union[str, Path]] = None
    ) -> bytes:
        """
        Generiert PDF aus HTML mit WeasyPrint
        """
        # Lazy import
        HTML, CSS = _try_import_weasyprint()
        
        if not HAS_WEASYPRINT or HTML is None:
            raise ImportError(
                f"WeasyPrint ist nicht verfügbar. "
                f"Auf Streamlit Cloud fehlen System-Bibliotheken (libpango). "
                f"Nutze stattdessen generate_pdf_reportlab(). "
                f"Fehler: {_weasyprint_error}"
            )
        
        # HTML rendern
        rendered_html = self.render_template_string(html_content, data)
        
        # CSS
        stylesheets = []
        if css:
            stylesheets.append(CSS(string=css))
        
        # PDF generieren
        html_doc = HTML(string=rendered_html)
        
        if output_path:
            html_doc.write_pdf(output_path, stylesheets=stylesheets)
            with open(output_path, 'rb') as f:
                return f.read()
        else:
            return html_doc.write_pdf(stylesheets=stylesheets)
    
    def generate_pdf_from_html_template(
        self,
        template_name: str,
        data: Dict[str, Any],
        css_file: Optional[str] = None,
        output_path: Optional[Union[str, Path]] = None
    ) -> bytes:
        """
        Generiert PDF aus HTML-Template-Datei
        """
        template_path = self.template_dir / template_name
        
        if not template_path.exists():
            raise FileNotFoundError(f"HTML-Template nicht gefunden: {template_path}")
        
        with open(template_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        css = None
        if css_file:
            css_path = self.template_dir / css_file
            if css_path.exists():
                with open(css_path, 'r', encoding='utf-8') as f:
                    css = f.read()
        
        return self.generate_pdf_from_html(html_content, data, css, output_path)
    
    def generate_pdf_reportlab(
        self,
        title: str,
        data: Dict[str, Any],
        sections: List[Dict[str, Any]],
        output_path: Optional[Union[str, Path]] = None,
        letterhead: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        Generiert PDF mit ReportLab (für komplexe Reports)
        """
        if not HAS_REPORTLAB:
            raise ImportError("reportlab ist nicht installiert: pip install reportlab")
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        styles = getSampleStyleSheet()
        
        # Eigene Styles
        styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=20
        ))
        styles.add(ParagraphStyle(
            name='FieldLabel',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=10
        ))
        styles.add(ParagraphStyle(
            name='FieldValue',
            parent=styles['Normal'],
            fontSize=10,
            leftIndent=20
        ))
        
        story = []
        
        # Titel
        story.append(Paragraph(title, styles['CustomTitle']))
        story.append(Paragraph(
            f"Erstellt am {datetime.now().strftime('%d.%m.%Y %H:%M')}",
            styles['Normal']
        ))
        story.append(Spacer(1, 20))
        
        # Sections verarbeiten
        for section in sections:
            section_type = section.get("type", "paragraph")
            
            if section_type == "heading":
                level = section.get("level", 1)
                text = self.render_template_string(section.get("text", ""), data)
                style_name = f'Heading{min(level, 6)}'
                story.append(Paragraph(text, styles[style_name]))
                story.append(Spacer(1, 10))
            
            elif section_type == "paragraph":
                text = self.render_template_string(section.get("text", ""), data)
                story.append(Paragraph(text, styles['Normal']))
                story.append(Spacer(1, 6))
            
            elif section_type == "field_value":
                label = section.get("label", "")
                field = section.get("field", "")
                value = data.get(field, "—")
                if isinstance(value, list):
                    value = ", ".join(str(v) for v in value)
                story.append(Paragraph(f"<b>{label}:</b> {value}", styles['Normal']))
                story.append(Spacer(1, 4))
            
            elif section_type == "table":
                headers = section.get("headers", [])
                rows = section.get("rows", [])
                
                table_data = [headers]
                for row in rows:
                    rendered_row = [
                        self.render_template_string(str(cell), data)
                        for cell in row
                    ]
                    table_data.append(rendered_row)
                
                table = Table(table_data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ]))
                story.append(table)
                story.append(Spacer(1, 10))
            
            elif section_type == "spacer":
                height = section.get("height", 20)
                story.append(Spacer(1, height))
        
        doc.build(story)
        
        pdf_bytes = buffer.getvalue()
        buffer.close()
        
        if output_path:
            with open(output_path, 'wb') as f:
                f.write(pdf_bytes)
        
        return pdf_bytes
    
    # ========================================
    # PDF Operationen
    # ========================================
    
    def merge_pdfs(
        self,
        pdf_files: List[Union[str, Path, bytes]],
        output_path: Optional[Union[str, Path]] = None
    ) -> bytes:
        """
        Fügt mehrere PDFs zusammen
        """
        if not HAS_PYPDF:
            raise ImportError("pypdf ist nicht installiert: pip install pypdf")
        
        merger = PdfMerger()
        
        for pdf in pdf_files:
            if isinstance(pdf, bytes):
                merger.append(io.BytesIO(pdf))
            else:
                merger.append(str(pdf))
        
        buffer = io.BytesIO()
        merger.write(buffer)
        merger.close()
        
        pdf_bytes = buffer.getvalue()
        
        if output_path:
            with open(output_path, 'wb') as f:
                f.write(pdf_bytes)
        
        return pdf_bytes
    
    def split_pdf(
        self,
        pdf_file: Union[str, Path, bytes],
        pages: List[int],
        output_path: Optional[Union[str, Path]] = None
    ) -> bytes:
        """
        Extrahiert bestimmte Seiten aus einem PDF
        """
        if not HAS_PYPDF:
            raise ImportError("pypdf ist nicht installiert: pip install pypdf")
        
        if isinstance(pdf_file, bytes):
            reader = PdfReader(io.BytesIO(pdf_file))
        else:
            reader = PdfReader(str(pdf_file))
        
        writer = PdfWriter()
        
        for page_num in pages:
            if 0 <= page_num < len(reader.pages):
                writer.add_page(reader.pages[page_num])
        
        buffer = io.BytesIO()
        writer.write(buffer)
        
        pdf_bytes = buffer.getvalue()
        
        if output_path:
            with open(output_path, 'wb') as f:
                f.write(pdf_bytes)
        
        return pdf_bytes
    
    def encrypt_pdf(
        self,
        pdf_file: Union[str, Path, bytes],
        password: str,
        output_path: Optional[Union[str, Path]] = None
    ) -> bytes:
        """
        Verschlüsselt ein PDF mit Passwort
        """
        if not HAS_PYPDF:
            raise ImportError("pypdf ist nicht installiert: pip install pypdf")
        
        if isinstance(pdf_file, bytes):
            reader = PdfReader(io.BytesIO(pdf_file))
        else:
            reader = PdfReader(str(pdf_file))
        
        writer = PdfWriter()
        
        for page in reader.pages:
            writer.add_page(page)
        
        writer.encrypt(password)
        
        buffer = io.BytesIO()
        writer.write(buffer)
        
        pdf_bytes = buffer.getvalue()
        
        if output_path:
            with open(output_path, 'wb') as f:
                f.write(pdf_bytes)
        
        return pdf_bytes
    
    # ========================================
    # Hilfsfunktionen
    # ========================================
    
    def get_document_hash(self, content: bytes) -> str:
        """Berechnet SHA-256 Hash eines Dokuments"""
        return hashlib.sha256(content).hexdigest()
    
    def get_available_features(self) -> Dict[str, bool]:
        """Gibt verfügbare Features zurück"""
        return {
            "docx_template": HAS_DOCXTPL,
            "docx_programmatic": HAS_DOCX,
            "pdf_html": HAS_WEASYPRINT,
            "pdf_reportlab": HAS_REPORTLAB,
            "pdf_operations": HAS_PYPDF,
        }


# ============================================
# Standard HTML/CSS für PDF
# ============================================

DEFAULT_PDF_CSS = """
@page {
    size: A4;
    margin: 2cm;
    @top-right {
        content: "Seite " counter(page) " von " counter(pages);
        font-size: 9pt;
        color: #666;
    }
}

body {
    font-family: 'Helvetica', 'Arial', sans-serif;
    font-size: 11pt;
    line-height: 1.5;
    color: #333;
}

h1 {
    font-size: 18pt;
    color: #1a1a1a;
    border-bottom: 2px solid #b45309;
    padding-bottom: 10px;
    margin-bottom: 20px;
}

h2 {
    font-size: 14pt;
    color: #333;
    margin-top: 20px;
}

h3 {
    font-size: 12pt;
    color: #555;
}

.field-group {
    margin-bottom: 15px;
}

.field-label {
    font-weight: bold;
    color: #555;
    font-size: 10pt;
}

.field-value {
    margin-left: 10px;
    padding: 5px 0;
}

table {
    width: 100%;
    border-collapse: collapse;
    margin: 15px 0;
}

th, td {
    border: 1px solid #ddd;
    padding: 8px;
    text-align: left;
}

th {
    background-color: #f5f5f5;
    font-weight: bold;
}

.header {
    display: flex;
    justify-content: space-between;
    align-items: center;
    margin-bottom: 30px;
    padding-bottom: 15px;
    border-bottom: 1px solid #eee;
}

.logo {
    max-height: 60px;
}

.meta-info {
    text-align: right;
    font-size: 9pt;
    color: #666;
}

.footer {
    margin-top: 30px;
    padding-top: 15px;
    border-top: 1px solid #eee;
    font-size: 9pt;
    color: #666;
}

.signature-area {
    margin-top: 50px;
    display: flex;
    justify-content: space-between;
}

.signature-box {
    width: 45%;
    border-top: 1px solid #333;
    padding-top: 5px;
    text-align: center;
    font-size: 9pt;
}
"""

DEFAULT_PDF_HTML = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{{ titel }}</title>
</head>
<body>
    <div class="header">
        <div class="logo-area">
            <h1>{{ kanzlei_name | default('RA-RHM Rechtsanwaltskanzlei') }}</h1>
        </div>
        <div class="meta-info">
            <p>Erstellt am: {{ datum }}<br>
            Dokument-Nr.: {{ dokument_nr | default('-') }}</p>
        </div>
    </div>
    
    <h1>{{ titel }}</h1>
    
    {% if beschreibung %}
    <p>{{ beschreibung }}</p>
    {% endif %}
    
    {% for section in sections %}
        {% if section.type == 'heading' %}
            <h{{ section.level | default(2) }}>{{ section.text }}</h{{ section.level | default(2) }}>
        
        {% elif section.type == 'paragraph' %}
            <p>{{ section.text }}</p>
        
        {% elif section.type == 'field_value' %}
            <div class="field-group">
                <span class="field-label">{{ section.label }}:</span>
                <span class="field-value">{{ data[section.field] | default_if_empty }}</span>
            </div>
        
        {% elif section.type == 'table' %}
            <table>
                <thead>
                    <tr>
                        {% for header in section.headers %}
                            <th>{{ header }}</th>
                        {% endfor %}
                    </tr>
                </thead>
                <tbody>
                    {% for row in section.rows %}
                        <tr>
                            {% for cell in row %}
                                <td>{{ cell }}</td>
                            {% endfor %}
                        </tr>
                    {% endfor %}
                </tbody>
            </table>
        
        {% elif section.type == 'spacer' %}
            <div style="height: {{ section.height | default(20) }}px;"></div>
        {% endif %}
    {% endfor %}
    
    {% if show_signature %}
    <div class="signature-area">
        <div class="signature-box">
            Ort, Datum
        </div>
        <div class="signature-box">
            Unterschrift
        </div>
    </div>
    {% endif %}
    
    <div class="footer">
        <p>{{ kanzlei_name | default('RA-RHM Rechtsanwaltskanzlei') }} | 
           {{ kanzlei_adresse | default('Adresse') }} | 
           Tel: {{ kanzlei_telefon | default('04331 732970') }} | 
           E-Mail: {{ kanzlei_email | default('info@ra-rhm.de') }}</p>
    </div>
</body>
</html>
"""


if __name__ == "__main__":
    # Test
    engine = DocumentEngine()
    print("Verfügbare Features:", engine.get_available_features())
